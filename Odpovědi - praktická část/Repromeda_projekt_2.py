from docx import Document
from docx.shared import Pt
import sys

def create_protocol(name, personal_number, sampling_date):
    # Vytvoření nového dokumentu Word
    doc = Document()

    # Přidání titulu
    title = doc.add_heading('Výsledný protokol genetického vyšetření', level=1)
    title.alignment = 1  # Zarovnání na střed

    # Přidání tabulky se zadanými informacemi
    table = doc.add_table(rows=3, cols=2)
    table.style = 'TableGrid'
    table.cell(0, 0).text = 'Jméno a příjmení:'
    table.cell(0, 1).text = name
    table.cell(1, 0).text = 'Rodné číslo:'
    table.cell(1, 1).text = personal_number
    table.cell(2, 0).text = 'Datum odběru:'
    table.cell(2, 1).text = sampling_date

    # Uložení dokumentu
    doc.save('genetic_protocol.docx')

if __name__ == "__main__":
    # Ověření, zda jsou poskytnuty všechny tři argumenty
    if len(sys.argv) != 4:
        print("Použití: python create_protocol.py <jméno> <rodné_číslo> <datum_odběru>")
    else:
        # Extrahování argumentů
        name = sys.argv[1]
        personal_number = sys.argv[2]
        sampling_date = sys.argv[3]

        # Vytvoření protokolu
        create_protocol(name, personal_number, sampling_date)
        print("Genetický protokol úspěšně vytvořen.")
